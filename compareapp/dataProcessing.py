import difflib
import fitz     # PDF reader
from django.conf import settings

# DOC generation
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfgen.canvas import Canvas
from datetime import datetime as date

# Excel importation
import pandas as pd

# Image importation
from PIL import Image as Img
import tensorflow as tf
import numpy as np


def compare_sections(section1, section2):      
    words1 = section1.strip().split()
    words2 = section2.strip().split()

    seq_matcher = difflib.SequenceMatcher(None, words1, words2)
    similarity = seq_matcher.ratio()
    is_different = similarity < 1.0

    added_words = []
    removed_words = []
    modified_words = []

    for tag, i1, i2, j1, j2 in seq_matcher.get_opcodes():
        if tag == 'replace':
            removed_words.append(' '.join(words1[i1:i2]))
            added_words.append(' '.join(words2[j1:j2]))
        elif tag == 'delete':
            removed_words.append(' '.join(words1[i1:i2]))
        elif tag == 'insert':
            added_words.append(' '.join(words2[j1:j2]))

    if similarity <= 0.4:
        modified_words.append(section2)

    if similarity == 1.0:
        tag = "S"  # Same
    elif added_words and not removed_words:
        tag = "A"  # Added
    elif removed_words and not added_words:
        tag = "R"  # Removed
    else:
        tag = "M"  # Modified       

    return similarity, is_different, tag, ' '.join(added_words), ' '.join(removed_words), ' '.join(modified_words) 

def read_docx(file_path):
    doc = Document(file_path)
    sections = {}
    current_section = None
    current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if text[0].isdigit() and (text[1] == '.' or (text[1].isdigit() and text[2] == '.')):
                if current_section:
                    sections[current_section] = '\n'.join(current_content)
                current_section = text
                current_content = []
            else:
                current_content.append(text)

    if current_section:
        sections[current_section] = '\n'.join(current_content)

    return sections

def read_pdf(file_path):
    doc = fitz.open(file_path)
    sections = {}
    current_section = None
    current_content = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text").strip()
        lines = text.splitlines()

        for line in lines:
            line = line.strip()
            if line:
                if line[0].isdigit() and (line[1] == '.' or (line[1].isdigit() and line[2] == '.')):
                    if current_section:
                        sections[current_section] = '\n'.join(current_content)
                    current_section = line
                    current_content = []
                else:
                    current_content.append(line)

    if current_section:
        sections[current_section] = '\n'.join(current_content)

    return sections

# DOC generation --------------------------

def create_report_docx(primary_data, data, reportNo, output_path, logo_path, comparedBy, short_description):
    new_doc = Document()

    # Set headers and footers for the entire document
    section = new_doc.sections[0]

    # Header with logo image
    header = section.header
    header_table = header.add_table(rows=3, cols=2, width=Inches(6))

    header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set table borders for header
    set_table_borders(header_table)

    # Set vertical alignment for all cells
    for row in header_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    header_table.cell(0, 0).paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.5))  # Adjust width as needed

    header_right_cell = header_table.cell(0, 1)
    header_right_para = header_right_cell.paragraphs[0]
    header_right_para.add_run("Documents Comparison Report").font.size = Pt(15)
    header_right_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    header_left_cell = header_table.cell(1, 0)
    header_left_para = header_left_cell.paragraphs[0]
    header_left_para.add_run("Compared By: " + comparedBy).font.size = Pt(12)
    header_left_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    header_right1_cell = header_table.cell(1, 1)
    header_right1_para = header_right1_cell.paragraphs[0]
    header_right1_para.add_run(f"Report Number: {reportNo}").font.size = Pt(12)
    header_right1_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    header_table.cell(2, 0).merge(header_table.cell(2, 1))

    header_right2_cell = header_table.cell(2, 0)
    header_right2_para = header_right2_cell.paragraphs[0]
    header_right2_para.add_run("Comparison short description: " + short_description).font.size = Pt(12)
    header_right2_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6))

    footer_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_table_borders(footer_table)

    for row in footer_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    footer_left_cell = footer_table.cell(0, 0)
    footer_left_para = footer_left_cell.paragraphs[0]
    cdate = date.now()
    footer_left_para.add_run(f"Comparison Date: {str(cdate).split(' ')[0]}").font.size = Pt(12)
    footer_left_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    footer_right_cell = footer_table.cell(0, 1)
    footer_right_para = footer_right_cell.paragraphs[0]
    footer_right_para.add_run("Page ").font.size = Pt(12)
    footer_right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    run = footer_right_para.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)

    run._r.append(OxmlElement('w:t'))

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

    run = footer_right_para.add_run(" of ")

    run = footer_right_para.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'NUMPAGES'
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)

    run._r.append(OxmlElement('w:t'))

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

    headers = set()
    for sections in data.values():
        headers.update(sections.keys())

    headers = sorted(headers, key=lambda x: (int(x.split('.')[0]), x))  # Sort headers

    if not primary_data:
        section = new_doc.sections[0]
        section.top_margin = Pt(350)
        section.bottom_margin = Pt(300)

        msg = new_doc.add_heading("Comparison Unsuccessful", level=1)
        msg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = msg.runs[0]
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(255, 0, 0)

        phrase = new_doc.add_paragraph("The document comparison has failed due to unsupported files content.")
        phrase.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    else:
        for header in headers:
            new_doc.add_heading(header, level=1)
            for doc_id, sections in data.items():
                section_content = sections.get(header, "")
                if section_content:
                    ref_section_content = list(data.values())[0].get(header, "")
                    similarity, is_different, tag , added_text, removed_text, modified_text = compare_sections(ref_section_content, section_content)
                    summary = "Same" if not is_different else "Different"
                    comparison_status = "Compared" if ref_section_content else "Not Compared"

                    if tag == 'A':
                        tag = 'Added'
                    elif tag == 'R':
                        tag = 'Removed'
                    elif tag == 'S':
                        tag = 'Similar'
                    else:
                        tag = 'Modified'

                    new_doc.add_heading(f"Document {doc_id}", level=2)

                    css = new_doc.add_paragraph()
                    run = css.add_run("Content Similarity Score: ")
                    run.bold = True
                    css.add_run(f"{int(similarity*100)}%")

                    s = new_doc.add_paragraph()
                    run = s.add_run("Summary: ")
                    run.bold = True
                    s.add_run(summary)
                    
                    t = new_doc.add_paragraph()
                    run = t.add_run("Tag: ")
                    run.bold = True
                    t.add_run(tag)

                    if modified_text:
                        new_doc.add_heading("Modified Text:", level=3)
                        new_doc.add_paragraph(modified_text)
                    else:
                        if added_text:
                            new_doc.add_heading("Added Text:", level=3)
                            new_doc.add_paragraph(added_text)
                        if removed_text:
                            new_doc.add_heading("Removed Text:", level=3)
                            new_doc.add_paragraph(removed_text)

                    new_doc.add_heading("Content:", level=3)
                    highlight_differences(new_doc, header, section_content, is_different)

    new_doc.save(output_path)

def set_table_borders(table):
    tbl = table._element
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)
    tbl.append(tbl_pr)

def highlight_differences(doc, title, text, is_different):
    paragraphs = text.split('\n')
    for para_text in paragraphs:
        para = doc.add_paragraph()
        for part in para_text.split(' '):
            run = para.add_run(part + ' ')
            if is_different:
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for differences

# PDF generation --------------------------

class NumberedCanvas(Canvas):
    def __init__(self, *args, **kwargs):
        Canvas.__init__(self, *args, **kwargs)
        self.pages = []

    def showPage(self):
        self.pages.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self.pages)
        for i, page in enumerate(self.pages):
            self.__dict__.update(page)
            self.draw_page_number(i + 1, num_pages)
            Canvas.showPage(self)
        Canvas.save(self)

    def draw_page_number(self, page_number, total_pages):
        self.setFont("Helvetica", 10)
        self.drawRightString(181 * mm, 14.5 * mm, f"Page {self.getPageNumber()} of {total_pages}")

def create_report_pdf(primary_data, data, report_no, output_path, logo_path, compared_by, short_description):
    pdf = SimpleDocTemplate(output_path, pagesize=A4, rightMargin=0.95 * inch, leftMargin=0.95 * inch, topMargin=0.75 * inch, bottomMargin=0.75 * inch)

    content = []

    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', fontSize=18, alignment=TA_CENTER, fontName='Helvetica-Bold')
    header_style = ParagraphStyle('Header', fontSize=12, alignment=TA_LEFT, textColor=colors.steelblue, fontName='Helvetica-Bold')
    normal_style = ParagraphStyle('Normal', fontSize=9, alignment=TA_LEFT)
    bold_style = ParagraphStyle('Bold', fontSize=9, alignment=TA_LEFT, fontName="Helvetica-Bold")

    # Header section
    if logo_path:
        logo_img = Image(logo_path, width=2 * inch, height=0.5 * inch)
    else:
        logo_img = ''

    header_data = [
        [logo_img, Paragraph("<b>Documents Comparison Report</b>", ParagraphStyle('Normal', fontSize=14, alignment=TA_CENTER, textColor=colors.darkcyan, fontName='Helvetica-Bold'))], 
        [Paragraph(f"<b>Compared By:</b> {compared_by}", ParagraphStyle('Normal', fontSize=8, alignment=TA_LEFT)), 
         Paragraph(f"<b>Report Number:</b> {report_no}", ParagraphStyle('Normal', fontSize=8, alignment=TA_RIGHT))],
        [Paragraph(f"<b>Short Description:</b> {short_description}",ParagraphStyle('Normal', fontSize=8, alignment=TA_LEFT)), ''] 
    ]
    
    header_table = Table(header_data, colWidths=[3.10 * inch, 3.10 * inch])
    header_table.setStyle(TableStyle([
        ('SPAN', (0, 2), (1, 2)),  # Merge short description columns [2, 0] to [2, 1]
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (1, 0), (1, 0), 'CENTER'),  # Center the "Document Comparison Report" text in [0, 1]
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),  # Add a border around the header
    ]))

    content.append(header_table)
    content.append(Spacer(1, 0.2 * inch))

    # Prepare headers for the content in sorted order
    headers = set()
    for sections in data.values():
        headers.update(sections.keys())
    headers = sorted(headers, key=lambda x: (int(x.split('.')[0]), x)) 

    # Loop through headers and add content
    if not primary_data:
        content.append(Spacer(1, 3 * inch))
        content.append(Paragraph("Comparison Unsuccessful", ParagraphStyle('Header', fontSize=18, alignment=TA_CENTER, textColor=colors.red, fontName='Helvetica-Bold')))

        content.append(Spacer(1, 0.3 * inch))
        content.append(Paragraph("The document comparion has failed due to unsupported file content.", ParagraphStyle('Normal', fontSize=12, alignment=TA_CENTER)))

    else:   
        for header in headers:
            content.append(Paragraph(header, header_style))
            content.append(Spacer(2, 0.1 * inch))

            for doc_id, sections in data.items():
                section_content = sections.get(header, "")
                if section_content:
                    ref_section_content = list(data.values())[0].get(header, "")
                    similarity, is_different, tag, added_text, removed_text, modified_text = compare_sections(ref_section_content, section_content)
                    summary = "Same" if not is_different else "Different"
                    tag_description = {
                        'A': 'Added',
                        'R': 'Removed',
                        'S': 'Similar',
                        'M': 'Modified'
                    }
                    tag = tag_description.get(tag, 'Unknown')

                    # Document section
                    content.append(Paragraph(f"<b>Document {doc_id}</b>", ParagraphStyle('Bold', fontSize=10, alignment=TA_LEFT, textColor=colors.slategray, fontName="Helvetica-Bold")))
                    content.append(Spacer(1, 0.1 * inch))

                    # Similarity score
                    content.append(Paragraph(f"<b>Content Similarity Score:</b> {int(similarity * 100)}%", normal_style))
                    content.append(Paragraph(f"<b>Summary:</b> {summary}", normal_style))
                    content.append(Paragraph(f"<b>Tag:</b> {tag}", normal_style))
                    content.append(Spacer(1, 0.1 * inch))

                    if modified_text:
                        content.append(Paragraph("<b>Modified Text:</b>", bold_style))
                        content.append(Paragraph(modified_text, ParagraphStyle('Normal', fontSize=9, textColor=colors.yellow, alignment=TA_LEFT)))
                    else:
                        if added_text:
                            content.append(Paragraph("<b>Added Text:</b>", bold_style))
                            content.append(Paragraph(added_text, ParagraphStyle('Normal', fontSize=9, textColor=colors.green, alignment=TA_LEFT)))
                        if removed_text:
                            content.append(Paragraph("<b>Removed Text:</b>", bold_style))
                            content.append(Paragraph(removed_text, ParagraphStyle('Normal', fontSize=9, textColor=colors.red, alignment=TA_LEFT)))

                    content.append(Paragraph("<b>Content:</b>", bold_style))
                    content.append(Paragraph(section_content, ParagraphStyle('Normal', fontSize=9, alignment=TA_LEFT)))

                    content.append(Spacer(1, 0.2 * inch))

    # Footer function for date and page number
    def add_page_footer(canvas, doc):
        canvas.saveState()
        width, height = A4

        # Footer text with page numbers in "Page X of Y" format
        footer_data = [
            [f"Comparison Date: {date.today().strftime('%Y-%m-%d')}", '']  # Page number X of Y
        ]

        footer_table = Table(footer_data, colWidths=[3.10 * inch, 3.10 * inch])
        footer_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),   # Align date to the left
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),  # Align page number to the right
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),  # Border around the footer
        ]))

        footer_table.wrapOn(canvas, width, height)
        footer_table.drawOn(canvas, inch, 0.5 * inch)

        canvas.restoreState()

    # Build PDF with footer on each page
    pdf.multiBuild(content, canvasmaker=NumberedCanvas, onFirstPage=add_page_footer, onLaterPages=add_page_footer)


# Excel comparison ------------------------

def compare_sheets(df1, df2):
    max_row = max(len(df1), len(df2))
    max_column = max(len(df1.columns), len(df2.columns))

    comparison_result = pd.DataFrame(index=range(max_row + 1), columns=range(max_column))

    unchanged_count = 0  # To count unchanged cells

    # Compare column headers
    for col in range(max_column):
        col_name1 = df1.columns[col] if col < len(df1.columns) else None
        col_name2 = df2.columns[col] if col < len(df2.columns) else None

        if col_name1 == col_name2:
            comparison_result.iloc[0, col] = col_name1  # No difference in headers
        else:
            comparison_result.iloc[0, col] = f'Modified: {col_name1} != {col_name2}'  # Mark header difference

    # Compare the data in the cells
    for row in range(max_row):
        for col in range(max_column):
            val1 = df1.iloc[row, col] if row < len(df1) and col < len(df1.columns) else None
            val2 = df2.iloc[row, col] if row < len(df2) and col < len(df2.columns) else None

            # Check if both values are NaN
            if pd.isna(val1) and pd.isna(val2):
                comparison_result.iloc[row + 1, col] = "-"  # Treat NaN as no change
                unchanged_count += 1
            elif val1 == val2:
                comparison_result.iloc[row + 1, col] = val1  # No difference
                unchanged_count += 1
            elif val1 is None:
                comparison_result.iloc[row + 1, col] = f'Added: {val2}'  # Added value
            elif val2 is None:
                comparison_result.iloc[row + 1, col] = f'Removed: {val1}'  # Removed value
            else:
                comparison_result.iloc[row + 1, col] = f'Modified: {val1} != {val2}'  # Different values

    similarity_score = (unchanged_count / (max_row * max_column)) * 100 if max_row * max_column > 0 else 0
    
    return comparison_result, similarity_score


# Image Processing ------------------------

def extract_text(img):
    try:
        # Load the image
        with open(img, 'rb') as image_file:
            image_bytes = image_file.read()
            text_response = settings.TEXTRACT_CLIENT.detect_document_text(
                Document={'Bytes': image_bytes}
            ) 
        extracted_text = ""
        for block in text_response['Blocks']:
            if block['BlockType'] == 'LINE':
                extracted_text += block['Text'] + " "
        
        return extracted_text
    
    except Exception as e:
        print(f"An unexpected error occurred during text extraction: {e}")
        return ""

def prepare_image(img):
    try:
        img = img.resize((224, 224))
        img_array = np.array(img)
        img_array = np.expand_dims(img_array, axis=0)
        img_array = tf.keras.applications.mobilenet_v2.preprocess_input(img_array)
    except Exception as e:
        print(f"An unexpected error occurred during image preparing: {e}")
    return img_array

def processImage(file):
    label = ""
    preScore = 0
    text = ""

    try:
        model = tf.keras.applications.mobilenet_v2.MobileNetV2(weights='imagenet')
        image = Img.open(file)
        processed_image = prepare_image(image)
        predictions = model.predict(processed_image)
        
        decoded_predictions = tf.keras.applications.mobilenet_v2.decode_predictions(predictions, top=1)[0]
        label = decoded_predictions[0][1]
        confidence = decoded_predictions[0][2]

        preScore = str(int(confidence))
        
    except Exception as e:
        print(f"An unexpected error occurred during image processing: {e}")
        
    text = extract_text(file)

    return {"label": label, "preScore": preScore, "text": text}
    

# end - Image Processing ------------------------