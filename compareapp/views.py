from django.conf import settings
from django.shortcuts import render, redirect, get_object_or_404, HttpResponse
from django.urls import reverse
from django.db.models import Q
from django.contrib.auth.decorators import login_required
from django.http import HttpRequest, JsonResponse

from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import user_passes_test

from django.views.decorators.csrf import csrf_exempt
from django.templatetags.static import static
from django.contrib import messages

from .forms import DocumentForm, CustomPasswordResetForm, UserForm, FeedbackForm, CustomSetPasswordForm
from .models import Document as Form, ComparisonReport, Feedback, UserLogs

from PyPDF2 import PdfReader
from docx import Document

import pandas as pd
from pathlib import Path
import json

import os
import openai
import requests
import logging
import datetime

# mail configuration  
from django.contrib.auth.tokens import default_token_generator
from django.utils.http import urlsafe_base64_encode
from django.utils.http import urlsafe_base64_decode
from django.core.mail import EmailMultiAlternatives
from django.utils.encoding import force_bytes
from django.template.loader import render_to_string
from django.contrib.auth.models import User

# Importing DOC and PDF generator
from .reportGenerator import create_report_docx, create_report_pdf, compare_sections, read_docx, read_pdf

def index(request):
    return render(request, "index.html")

@login_required
def logoutUser(request):

    log = UserLogs.objects.create(
        user = request.user,
        done_by = request.user.get_full_name() or request.user.username,
        last_login = request.user.last_login,
        action = "Logged out",
        action_type = "logout"
    )

    log.save()
    
    logout(request)
    messages.info(request, "You have been logged out.")
    return redirect('index')

def loginUser(request):
    if request.method == "POST":
        username = request.POST.get('userid')
        password = request.POST.get('password')

        if not username or not password:
            messages.warning(request, "Please provide the login credentials.")
        else:
            user = authenticate(request, username=username, password=password)
            if user is not None:
                login(request, user)
                
                request.session.set_expiry(settings.SESSION_COOKIE_AGE)
                request.session['expiry_time'] = settings.SESSION_COOKIE_AGE

                log = UserLogs.objects.create(
                    user = request.user,
                    done_by = request.user.get_full_name() or request.user.username,
                    last_login = request.user.last_login,
                    action = "Logged In",
                    action_type = "login"
                )

                log.save()

                messages.success(request, "You have successfully logged in.")
                return redirect('dashboard')
            else:
                messages.error(request, "Please provide valid login credentials.")
            
    return render(request, "login.html")

@login_required
def submitFeedback(request):
    previous_url = request.META.get('HTTP_REFERER', 'dashboard')

    if request.method == 'POST':
        form = FeedbackForm(request.POST)
        if form.is_valid():
            form.save()

            log = UserLogs.objects.create(
                user = request.user,
                done_by = request.user.get_full_name() or request.user.username,
                last_login = request.user.last_login,
                action = "Submit Feedback",
                action_type = "create"
            )

            log.save()

            messages.success(request, 'Thank you for your valuable feedback!')
        else:
            messages.warning(request, 'Please provide valid feedback!')

    return redirect(previous_url)
      
# User Management Section -------------------------------------------------------------
@login_required
@user_passes_test(lambda user: user.is_superuser)
def userManagement(request):   
    query = request.GET.get('q')
    filter_by = request.GET.get('status')

    users = User.objects.all()

    if query:
        users = users.filter(
            Q(id__icontains=query) |
            Q(username__icontains=query) |
            Q(email__icontains=query)
        ).distinct()
    
    if filter_by:
        if filter_by == 'active':
            users = users.filter(is_active=True)
        elif filter_by == 'inactive':
            users = users.filter(is_active=False)

    return render(request, 'user-management/users.html', {'users': users})

@login_required
@user_passes_test(lambda user: user.is_superuser)
def userLogs(request: HttpRequest):
    query = request.GET.get('q')
    filter_by = request.GET.get('status')
    
    user_id = request.GET.get("user_id", '')
    
    if user_id:
        logs = UserLogs.objects.filter(user=user_id)
    else:
        logs = UserLogs.objects.all()

    if query:
        logs = logs.filter(
            Q(action__icontains=query) |
            Q(done_by__icontains=query)
        ).distinct()

    if filter_by:
        logs = logs.filter(action_type=filter_by)

    return render(request, 'user-management/user-logs.html', {'logs': logs[::-1]})

@login_required
@user_passes_test(lambda user: user.is_superuser)
def add_edit_user(request, user_id=None):
    if user_id:
        user = get_object_or_404(User, id=user_id)
        form = UserForm(request.POST or None, request.FILES or None , instance=user)
        user_permissions = user.user_permissions.all()
    else:
        user = None
        form = UserForm(request.POST or None, request.FILES or None , request=request)
        user_permissions = []

    if request.method == 'POST':
        if form.is_valid():
            user = form.save()
            user.user_permissions.set(form.cleaned_data['permissions'])

            if not user_id:
                passwd_type = form.cleaned_data.get('password_type')
                if passwd_type == "bymail":
                    # Send email for password creation
                    subject = "Create Your Password"
                    email_template_name = "password-base/password_creation_email.html"
                    context = {
                        "email": user.email,
                        "domain": request.META['HTTP_HOST'],
                        "site_name": "Doc Comparison Pro",
                        "uid": urlsafe_base64_encode(force_bytes(user.pk)),
                        "user": user,
                        "token": default_token_generator.make_token(user),
                        "protocol": "http",
                    }
                    email_message = render_to_string(email_template_name, context)
                    email = EmailMultiAlternatives(
                        subject=subject,
                        body=email_message,
                        from_email=settings.DEFAULT_FROM_EMAIL,
                        to=[user.email]
                    )
                    email.attach_alternative(email_message, "text/html")
                    email.send(fail_silently=False)

                    messages.success(request, 'User created successfully.')
                    messages.info(request, 'Password creation request has been sent to the user.')
                else:
                    messages.success(request, 'User created successfully.')
                
                log = UserLogs.objects.create(
                    user = request.user,
                    done_by = request.user.get_full_name() or request.user.username,
                    last_login = request.user.last_login,
                    action = "User Created",
                    action_type = "create"
                )

                log.save()
                    
            else:
                messages.success(request, 'User updated successfully.')

            log = UserLogs.objects.create(
                user = request.user,
                done_by = request.user.get_full_name() or request.user.username,
                last_login = request.user.last_login,
                action = "User Updated",
                action_type = "update"
            )

            log.save()

            return redirect('user-management')

    return render(request, 'user-management/user_form.html', {
        'form': form,
        'user': user,
        'user_permissions': user_permissions,
    })
    
@login_required
@user_passes_test(lambda user: user.is_superuser)
def delete_user(request, user_id):
    user = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        user.delete()
        messages.success(request, f'User {user.username} was successfully deleted.')

        log = UserLogs.objects.create(
            user = request.user,
            done_by = request.user.get_full_name() or request.user.username,
            last_login = request.user.last_login,
            action = "User Removed",
            action_type = "delete"
        )

        log.save()
        
        return redirect(reverse('user-management'))
    return render(request, 'user-management/user_management.html', {'users': User.objects.all()})

# User Profile View
@login_required
def user_profile(request, user_id):
    exists = User.objects.filter(id=user_id)

    if not exists:
        messages.warning(request, "The requested profile is not available!")
        return redirect(request.META.get('HTTP_REFERER', 'dashboard'))
    
    if request.user.is_superuser:
        user = User.objects.get(id=user_id)
    else:
        user = User.objects.get(id=request.user.id)

    total_comparison = len(ComparisonReport.objects.filter(user=request.user))
    total_documents = len(Form.objects.filter(user=request.user))
    failed_comparisons = len(ComparisonReport.objects.filter(user=request.user, comparison_status=False))
    activities = UserLogs.objects.filter(user=request.user).last()
    last_activity = activities.action if activities else ""

    if not request.session.get(f"viewed_up_{user_id}_{request.user}"):
        log = UserLogs.objects.create(
            user = request.user,
            done_by = request.user.get_full_name() or request.user.username,
            last_login = request.user.last_login,
            action = "Viewed Profile",
            action_type = "open"
        )
        
        log.save()

        request.session[f"viewed_up_{user_id}_{request.user}"] = True

    specific_permissions = [
        "auth.add_user",
        "auth.change_user",
        "auth.delete_user",
        "auth.view_user",
        "compareapp.add_comparisonreport",
        "compareapp.change_comparisonreport",
        "compareapp.delete_comparisonreport",
        "compareapp.view_comparisonreport",
        "compareapp.add_document",
        "compareapp.change_document",
        "compareapp.delete_document",
        "compareapp.view_document",
        "compareapp.add_feedback",
        "compareapp.change_feedback",
        "compareapp.delete_feedback",
        "compareapp.view_feedback",
    ]

    return render(request, 'user-management/user_profile.html', {
        'user': user,
        'total_comparison': total_comparison,
        'total_documents': total_documents,
        'failed_comparisons': failed_comparisons,
        'specific_permissions': specific_permissions,
        'last_activity': last_activity,
    })

# Comparison analytics ------------------------------------
@login_required
def analytics(request):
    if not request.user.is_authenticated:
        messages.warning(request, "Login Required!")
        return redirect('login')

    files_format = ['pdf', 'docx', 'xlsx', 'pptx', 'vsd', 'wav', 'mp4', 'png', 'txt', 'other']
    file_labels = ['PDFs', 'Documents', 'Spreadsheets', 'Prasentations', 'Visios', 'Audios', 'Videos', 'Images', 'Text', 'Others']

    # Handling users
    if request.user.is_superuser:
        total_files_data = [ len(Form.objects.filter(comparison_between = doc)) for doc in files_format ]
        all_comparison_data = [ len(ComparisonReport.objects.filter(comparison_between = doc)) for doc in files_format ]

        total_comparisons = len(ComparisonReport.objects.all())
        failed_reports = len(ComparisonReport.objects.filter(comparison_status=False))
        success_reports = len(ComparisonReport.objects.filter(comparison_status=True))
    else:
        total_files_data = [ len(Form.objects.filter(user=request.user, comparison_between = doc)) for doc in files_format ]
        all_comparison_data = [ len(ComparisonReport.objects.filter(user=request.user, comparison_between = doc)) for doc in files_format ]

        total_comparisons = len(ComparisonReport.objects.filter(user=request.user))
        failed_reports = len(ComparisonReport.objects.filter(user=request.user, comparison_status=False))
        success_reports = len(ComparisonReport.objects.filter(user=request.user, comparison_status=True))
    
    total_users = len(User.objects.all())
    user_feedbacks = len(Feedback.objects.all())

    total_files = {
        'labels': file_labels,
        'values': total_files_data
    }

    all_comparisons = {
        'labels': file_labels,
        'values': all_comparison_data
    }
    
    report_data={
        'labels': file_labels,
        'values': total_files_data
    }

    if not request.session.get(f"opened_analytics_{request.user.username}"):
        log = UserLogs.objects.create(
            user = request.user,
            done_by = request.user.get_full_name() or request.user.username,
            last_login = request.user.last_login,
            action = "Opened Analytics",
            action_type = "open"
        )

        log.save()

        request.session[f"opened_analytics_{request.user.username}"] = True

    return render(request, 'analytics.html', { 'failed_reports': failed_reports, "success_reports": success_reports, 'total_users': total_users, 'total_files': total_files, 'report_data':report_data , 'user_feedbacks': user_feedbacks, 'total_comparisons':total_comparisons , 'all_comparisons':all_comparisons})

def password_reset_request(request):
    if request.method == "POST":
        form = CustomPasswordResetForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['email']
            users = User.objects.filter(email=email)
            if users.exists():
                for user in users:
                    subject = "Password Reset Requested"
                    email_template_name = "password-base/password_reset_email.html"
                    context = {
                        "email": user.email,
                        "domain": request.META['HTTP_HOST'],
                        "site_name": "Doc Comparison Pro",
                        "uid": urlsafe_base64_encode(force_bytes(user.pk)),
                        "user": user,
                        "token": default_token_generator.make_token(user),
                        "protocol": "http",
                    }
                    email_message = render_to_string(email_template_name, context)
                    email = EmailMultiAlternatives(
                        subject=subject,
                        body=email_message,
                        from_email=settings.DEFAULT_FROM_EMAIL,
                        to=[user.email]
                    )
                    email.attach_alternative(email_message, "text/html")
                    email.send(fail_silently=False)
                messages.success(request, "Password reset request has been sent.")

                log = UserLogs.objects.create(
                    user = request.user,
                    done_by = request.user.get_full_name() or request.user.username,
                    last_login = request.user.last_login,
                    action = "Password Reset, requested",
                    action_type = "update"
                )

                log.save()

            else:
                messages.error(request, "The provided email is not registered.")
            form = CustomPasswordResetForm()
    else:
        form = CustomPasswordResetForm()

    return render(request, "password-base/password_reset_form.html", {"form": form})

def password_creation_view(request, uidb64, token):
    try:
        uid = urlsafe_base64_decode(uidb64).decode()
        user = get_object_or_404(User, pk=uid)
    except (TypeError, ValueError, OverflowError, User.DoesNotExist):
        user = None

    if user is not None and default_token_generator.check_token(user, token):
        if request.method == 'POST':
            form = CustomSetPasswordForm(user, request.POST)
            if form.is_valid():
                form.save()

                log = UserLogs.objects.create(
                    user = request.user,
                    done_by = request.user.get_full_name() or request.user.username,
                    last_login = request.user.last_login,
                    action = "Password Created",
                    action_type = "create"
                )

                log.save()                
                
                return redirect('password_create_done')
        else:
            form = CustomSetPasswordForm(user)
    else:
        form = None

    return render(request, 'password-base/password_creation_confirm.html', {'form': form})

@login_required
def dashboard(request):
    if not request.user.is_authenticated:
        messages.warning(request, "Login Required!")
        return redirect('login')
    
    query = request.GET.get('q', '')
    filter_by = request.GET.get('filter')

    if request.user.is_superuser:
        reports = ComparisonReport.objects.all()
    else:
        reports = ComparisonReport.objects.filter(user=request.user)
    
    if filter_by:
        valid_filters = ['docx', 'pdf', 'xlsx', 'pptx', 'vsd', 'wav', 'mp4', 'png', 'txt', 'other']
        if filter_by in valid_filters:
            reports = reports.filter(comparison_between__icontains=filter_by)

    if query:
        reports = reports.filter(
            Q(report_number__icontains=query) |
            Q(short_description__icontains=query) |
            Q(department_type__icontains=query) |
            Q(comparison_date__icontains=query) |
            Q(compared_by__icontains=query)
        ).distinct()
    
    return render(request, 'dashboard.html', { "reports": reports[::-1] })

@login_required
def viewComparison(request, report_id):
    if not request.user.is_authenticated:
        messages.warning(request, "Login Required!")
        return redirect('login')
    
    try:
        if request.user.is_superuser:
            report = ComparisonReport.objects.filter(report_number=report_id).first()
            compared_documents = report.compared_documents
            comparison_status = report.comparison_status
            document_ids = list(compared_documents.values())
            documents = Form.objects.filter(document_id__in=document_ids)
        else:
            report = ComparisonReport.objects.filter(user=request.user, report_number=report_id).first()
            compared_documents = report.compared_documents
            comparison_status = report.comparison_status
            document_ids = list(compared_documents.values())
            documents = Form.objects.filter(user=request.user, document_id__in=document_ids)
        
        if not request.session.get(f"viewed_comparison_{report_id}"):
            log = UserLogs.objects.create(
                user = request.user,
                done_by = request.user.get_full_name() or request.user.username,
                last_login = request.user.last_login,
                action = f"Viewed comparison info, RN-{report_id}",
                action_type = "read"
            )

            log.save()

            request.session[f"viewed_comparison_{report_id}"] = True

    except:
        messages.warning(request, "The requested report is not available!")
        return redirect('dashboard')

    comparison_details = report.comparison_summary

    return render(request, "view-comparisons/view-comparison.html", {
        "documents": documents,
        "comparison_status": comparison_status,
        "comparison_details": comparison_details,
        "report": report_id,
        "report_summary": report.ai_summary
    })

@login_required
def formView(request):   
    if not request.user.is_authenticated:
        messages.warning(request, "Login Required!")
        return redirect('login')

    report_number = request.GET.get('report_number')
    success = request.GET.get('success')
    saved = True if (request.GET.get('saved', False) == 'True') else False
    last_report = ComparisonReport.objects.last()
    current_date = datetime.date.today()

    if last_report:
        new_report_number = f"DC{ request.user.id }R{int(last_report.report_number.split('R')[1]) + 1}"
    else:
        new_report_number = f"DC{ request.user.id }R1001"

    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        valid_format = ['docx', 'pdf']

        if form.is_valid():
            saveReportNumber = request.POST.get("report_number")
            comparison_between = request.POST.get("documents_format")
            comparison_date = request.POST.get("comparison_date")
            short_description = request.POST.get("short_description")
            description = request.POST.get("description")
            department_type = request.POST.get("department_type")
            documents = request.FILES.getlist('upload_documents')

            if not comparison_between:
                messages.warning(request, f"Please select the files type or reset to intialise the upload process!")
                return redirect('form')

            elif comparison_between not in valid_format:
                messages.warning(request, f"Bad request, invalid documents format '{comparison_between}' for comaprison!")
                return redirect('form')

            for doc in documents:
                file_extension = os.path.splitext(doc.name)[1].lstrip('.').lower()

                if file_extension != comparison_between:
                    messages.error(request, f"Bad request, file extension '{file_extension}' mismatched with selected format '{comparison_between}'.")
                    return redirect('form')
            
            try:
                for doc in documents:
                    document_instance = Form()
                    document_instance.comparison_between = comparison_between
                    document_instance.upload_documents = doc
                    document_instance.user = request.user
                    document_instance.save()

                comparison_instance = ComparisonReport()
                comparison_instance.report_number = saveReportNumber
                comparison_instance.department_type = department_type
                comparison_instance.description = description
                comparison_instance.short_description = short_description
                comparison_instance.comparison_date = comparison_date
                comparison_instance.comparison_between = comparison_between
                comparison_instance.user = request.user
                comparison_instance.save()
            except:
                messages.warning(request, "Error occured while saving the files!")
                messages.info(request, "Please reset the upload process!")
                return redirect("form")

            url = reverse('form')
            redirect_url = f"{url}?saved=True&report_number={saveReportNumber}"

            log = UserLogs.objects.create(
                user = request.user,
                done_by = request.user.get_full_name() or request.user.username,
                last_login = request.user.last_login,
                action = "Documents uploaded",
                action_type = "create"
            )

            log.save()
            
            return redirect(redirect_url)
        else:
            messages.warning(request, "All fields are required to be fill!")
    else:
        form = DocumentForm()

    return render(request, "form.html", {
        'form': form,
        'success': success,
        'saved': saved,
        'report_number': report_number,
        'current_date': current_date,
        'new_report_number': new_report_number
    })
    
# Import data view
def importData(request):
    if request.method == 'POST' and request.FILES.get('excelFile'):
        try:
            excel_file = request.FILES['excelFile']
            df = pd.read_excel(excel_file)

            description = df['description'].iloc[0]
            short_description = df['short_description'].iloc[0]
            department_type = df['department_type'].iloc[0]

            return JsonResponse({
                'success': True,
                'data': {
                    'description': description,
                    'short_description': short_description,
                    'department_type': department_type,
                }
            })

        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

# Comparison Cancellation Process
@login_required
def cancelComparison(request: HttpRequest):

    report_number = request.GET.get('report_number', '')
    documents = Form.objects.filter(new=True, user=request.user)
    report = ComparisonReport.objects.filter(report_number=report_number, user=request.user)

    if report:
        report.delete()
        
    if documents:
        documents.delete()

    messages.success(request, "The comparison was cancelled, resetting the process.")

    log = UserLogs.objects.create(
        user = request.user,
        done_by = request.user.get_full_name() or request.user.username,
        last_login = request.user.last_login,
        action = "Comparison Cancelled",
        action_type = "delete"
    )

    log.save()

    return redirect('form')

@login_required
def documentDetail(request, doc_id):
    if not request.user.is_authenticated:
        messages.warning(request, "Login Required!")
        return redirect('login')
    
    try: 
        if request.user.is_superuser:
            document = get_object_or_404(Form, document_id=doc_id)
        else:
            document = get_object_or_404(Form, document_id=doc_id, user=request.user)
    except:
        messages.warning(request, "Invalid document ID, please provide valid ID.")
        return redirect('dashboard')
    
    if not request.session.get(f"opened_d_{document.report_number}"):
        log = UserLogs.objects.create(
            user = request.user,
            done_by = request.user.get_full_name() or request.user.username,
            last_login = request.user.last_login,
            action = f"Opened compared document, DN-{document.report_number}/{document.document_id}",
            action_type = "open"
        )

        log.save()

        request.session[f"opened_d_{document.report_number}"] = True

    return render(request, 'document-details.html', { 'document': document })

@login_required
def comparedDocument(request, id):
    if not request.user.is_authenticated:
        messages.warning(request, "Login Required!")
        return redirect('login')

    query = request.GET.get('q', '')
    filter_type = request.GET.get('filter', '')

    try:
        documents = Form.objects.filter(report_number=id, user=request.user)
    except:
        messages.warning(request, "Document ID is not available or invalid!")
        return redirect('dashboard')

    if not request.session.get(f"viewed_cd_{id}"):
        log = UserLogs.objects.create(
            user = request.user,
            done_by = request.user.get_full_name() or request.user.username,
            last_login = request.user.last_login,
            action = f"Viewed compared documents, RN-{id}",
            action_type = "read"
        )

        log.save()

        request.session[f"viewed_cd_{id}"] = True

    if query:
        documents = documents.filter(
            Q(document_id__icontains=query) |
            Q(comparison_between__icontains=query) |
            Q(report_number__icontains=query) |
            Q(summary__icontains=query)
        ).distinct()

    if filter_type and filter_type != '':
        documents = documents.filter(comparison_between=filter_type)
    
    return render(request, 'view-comparisons/compared-documents.html', { 'documents': documents })

# Optional route for future use --------
@login_required
def removeDocument(request, doc_id):
    previous_url = request.META.get('HTTP_REFERER', 'dashboard')

    try:
        document = get_object_or_404(Form, document_id=doc_id)
        document.delete()
        remove_file = document.upload_document.path
        os.remove(remove_file)
        messages.success(request, "Document deleted successfully.")

        log = UserLogs.objects.create(
            user = request.user,
            done_by = request.user.get_full_name() or request.user.username,
            last_login = request.user.last_login,
            action = "Removed Document",
            action_type = "delete"
        )

        log.save()

    except:
        messages.warning(request, "Invalid document ID, please provide valid ID")
        
    return redirect(previous_url)

@login_required
def comparison(request: HttpRequest):
    if not request.user.is_authenticated:
        messages.warning(request, "Login Required!")
        return redirect('login')

    old_report_number = request.GET.get('report_number', '')
    documents = Form.objects.filter(new=True, user=request.user)

    if not documents or len(documents) < 2:
        messages.warning(request, "Minimum two documents are required to perform the comparison!")
        return redirect('form') 
    
    comparison_between = documents[0].comparison_between
    user_full_name = request.user.get_full_name().title()
    
    try:
        short_description = get_object_or_404(ComparisonReport, report_number=old_report_number).short_description
    except:
        messages.error(request, "Invalid comparison ID, can't perform the comparison!")
        return redirect("form")
    
    if not user_full_name:
        comparedBy = request.user.username.title()
    else:
        comparedBy = user_full_name  
    
    data = {}
    ai_summary = {}
    for doc in documents:
        file_path = doc.upload_documents.path
        if comparison_between == 'docx':
            sections = read_docx(file_path)
            data[doc.document_id] = sections
            
            # Getting AI Summary
            content = read_file(file_path)
            ai_summary[doc.document_id] = getSummary(content)
            
        elif comparison_between == 'pdf':
            sections = read_pdf(file_path)
            data[doc.document_id] = sections
            
            # Getting AI Summary
            content = read_file(file_path)
            ai_summary[doc.document_id] = getSummary(content)

        else:
            messages.error(request, "Can't perform comparison due to invalid file format.")
            return redirect('form')

    result_dir = os.path.join(settings.MEDIA_ROOT, 'comparison-reports')
    os.makedirs(result_dir, exist_ok=True)

    docx_path = os.path.join(result_dir, f"{old_report_number}.docx")
    pdf_path = os.path.join(result_dir, f"{old_report_number}.pdf")
    logo_path = "compareapp" + static('images/logo.png')
    primary_data = data[documents[0].document_id] or ""

    # Generating reports
    create_report_docx(primary_data, data, old_report_number, docx_path, logo_path, comparedBy, short_description)
    create_report_pdf(primary_data, data, old_report_number, pdf_path, logo_path, comparedBy, short_description)

    # Prepare comparison details
    comparison_details = {}
    overall_similarity_scores = {}
    headers = set()
    for sections in data.values():
        headers.update(sections.keys())

    headers = sorted(headers, key=lambda x: (int(x.split('.')[0]), x))  # Sort headers

    primary_doc_id = list(data.keys())[0]  # Assume the first document is the primary document
    for header in headers:
        comparison_details[header] = {
            'primary': data[primary_doc_id].get(header, ""),
            'documents': {}
        }
        ref_section_content = data[primary_doc_id].get(header, "")
        for doc_id, sections in data.items():
            if doc_id != primary_doc_id:
                section_content = sections.get(header, "")
                if section_content:
                    similarity, is_different, tag ,added_text, removed_text, modified_text  = compare_sections(ref_section_content, section_content)
                    if is_different:  # Only include if different
                        comparison_details[header]['documents'][doc_id] = {
                            'content': section_content or 'Removed',
                            'tag': tag,
                            'added_text': added_text,
                            'removed_text': removed_text,
                            'modified_text': modified_text
                        }
                    else:
                        comparison_details[header]['documents'][doc_id] = {
                            'content': 'Same as Primary Document',
                            'tag': 'S'
                        }

    # Calculate overall similarity score for each document
    ref_doc_content = "\n".join(list(data.values())[0].values())
    for doc_id, sections in data.items():
        doc_content = "\n".join(sections.values())
        overall_similarity_score, _, _, _, _, _ = compare_sections(ref_doc_content, doc_content )
        overall_similarity_scores[doc_id] = int(overall_similarity_score * 100)


    # Now Saving the Comparison Result
    for doc in documents:
        if not data[primary_doc_id] or not data[doc.document_id]:
            doc.summary = "Not Applicable"
            doc.similarity_score = "Not Applicable"
            doc.comparison_status = 'Not Compared'
        else:
            doc.summary = "Same" if overall_similarity_scores[doc.document_id] == 100 else "Different"
            doc.similarity_score = f"{int(overall_similarity_scores[doc.document_id])}%"
            doc.comparison_status = 'Compared'

        doc.new = False
        doc.ai_summary = ai_summary[doc.document_id]
        doc.report_number = old_report_number
        doc.save()

    compared_documents = {}
    for index, doc in zip(range(1, len(documents) + 1), documents):
        compared_documents[f'doc{index}'] = doc.document_id

    prepare_data = read_file(docx_path)

    if not comparison_details or not data[primary_doc_id]:
        comparison_status = False
        comparison_ai_summary = "The document comparison has failed due to unsupported documents, but you can still get an AI-generated summary of the provided documents."
    else:
        comparison_status = True
        comparison_ai_summary = getSummary(prepare_data, ind=False)

    try:
        comparison_instance = ComparisonReport.objects.get(report_number=old_report_number, user=request.user)
        comparison_instance.compared_documents = compared_documents
        comparison_instance.comparison_summary = comparison_details
        comparison_instance.ai_summary = comparison_ai_summary
        comparison_instance.comparison_status = comparison_status
        comparison_instance.compared_by = comparedBy
        comparison_instance.report_path = docx_path
        comparison_instance.save()

        log = UserLogs.objects.create(
            user = request.user,
            done_by = request.user.get_full_name() or request.user.username,
            last_login = request.user.last_login,
            action = "Compared, Documents",
            action_type = "create"
        )

        log.save()

        return redirect(f'{reverse("form")}?success=True&report_number={old_report_number}')

    except Exception as e:
        # messages.error(request, "Error occured while saving the comparison data.")
        return HttpResponse(f"Error: {e}")

@login_required
def preview(request, report):
    comparison_report = Path(settings.MEDIA_ROOT) / f'comparison-reports/{report}.docx'
    pdf_path = comparison_report.with_suffix('.pdf')
    pdf_url = str(pdf_path.relative_to(settings.MEDIA_ROOT)).replace("\\", "/")
    
    if not request.session.get(f"opened_cr_{report}"):
        log = UserLogs.objects.create(
            user = request.user,
            done_by = request.user.get_full_name() or request.user.username,
            last_login = request.user.last_login,
            action = f"Opened Comparison Report, RN-{report}",
            action_type = "open"
        )

        log.save()

        request.session[f"opened_cr_{report}"] = True
        
    return render(request, 'report-preview.html', {'pdf_path': f'/media/{pdf_url}', 'report': report})

def softwareDocumentation(request):

    if not request.session.get(f"viewed_sd_{request.user}"):
        log = UserLogs.objects.create(
            user = request.user,
            done_by = request.user.get_full_name() or request.user.username,
            last_login = request.user.last_login,
            action = "Viewed Documentation",
            action_type = "open"
        )

        log.save()

        request.session[f"viewed_sd_{request.user}"] = True
    
    return render(request, "documentation/view.html")

# Chatting with document -----------------------------------------------------------------------
@csrf_exempt
def uploadPDF(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            relative_pdf_path = data.get('file_url')
            try:
                RN = relative_pdf_path.split('/')[-1].split('.')[0]
            except:
                RN = ''

            absolute_pdf_path = os.path.join(settings.MEDIA_ROOT, relative_pdf_path)

            if not os.path.exists(absolute_pdf_path):
                return JsonResponse({'error': 'File not found or invalid file path'}, status=400)

            with open(absolute_pdf_path, 'rb') as pdf_file:
                files = [
                    ('file', ('file', pdf_file, 'application/pdf'))
                ]
                headers = {
                    'x-api-key': settings.CHATPDF_API_KEY
                }

                response = requests.post('https://api.chatpdf.com/v1/sources/add-file', headers=headers, files=files)

                if response.status_code == 200:

                    log = UserLogs.objects.create(
                        user = request.user,
                        done_by = request.user.get_full_name() or request.user.username,
                        last_login = request.user.last_login,
                        action = f"Chat with comparison report, RN-{RN}",
                        action_type = "read"
                    )

                    log.save()

                    return JsonResponse(response.json())
                else:
                    return JsonResponse({'error': response.text}, status=response.status_code)
        except Exception as e:
            print(f"Exception: {str(e)}")
            return JsonResponse({'error': str(e)}, status=500)
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=400)
    
logger = logging.getLogger(__name__)

@csrf_exempt
def proxy_chat_pdf(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            question = data.get('question')
            source_id = data.get('source_id')
            headers = {
                'x-api-key': settings.CHATPDF_API_KEY,
                'Content-Type': 'application/json'
            }
            payload = {
                'sourceId': source_id,
                'messages': [
                    {'role': 'user', 'content': question}
                ]
            }
            response = requests.post('https://api.chatpdf.com/v1/chats/message', json=payload, headers=headers)
            
            if response.ok:
                return JsonResponse(response.json())
            else:
                logger.error(f"ChatPDF API Error: {response.status_code} - {response.text}")
                return JsonResponse(response.json(), status=response.status_code)
        except Exception as e:
            logger.exception("An error occurred while processing the request")
            return JsonResponse({'error': str(e)}, status=500)
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=400)

        
# Handling an environment for generating comparison summary --------------------------------------
def read_word_data(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def read_pdf_data(file):
    pdf = PdfReader(file)
    content = ""
    for page in pdf.pages:
        content += page.extract_text()
    return content

def read_file(file):
    file_type = os.path.splitext(file)[1].lower()
    if file_type == ".docx":
        return read_word_data(file)
    elif file_type == ".pdf":
        return read_pdf_data(file)
    else:
        return None

def getSummary(data, ind=True):
    openai.api_key = settings.OPENAI_API_KEY

    try:
        if ind:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an AI that give the summary of document content."},
                    {"role": "user", "content": f"give the summary of the following document content:\n{data}"}
                ],
                max_tokens=2000
            )

            summary = response['choices'][0]['message']['content'].strip()
            return summary
        else:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                {"role": "system", "content": "You are an AI that compares documents and highlights revisions/changes."},
                {"role": "user", "content": f"Compare the following documents and highlight the revisions/changes:\n{data}"}
            ],
                max_tokens=2000
            )

            summary = response['choices'][0]['message']['content'].strip()

            return summary
    except Exception as e:
        print("Error occured while fetching the summary response!\n", e)

        return None
    
# end - Handling an environment for generating comparison summary --------------------------------------
